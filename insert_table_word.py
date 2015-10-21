#coding:utf-8
from docx import Document
def inert_table(file_name_path,cols_int,title_list,info_list,replace_str):
    count = 0
    document = Document(file_name_path)
    for paragraph  in document.paragraphs:
        count += 1
        if paragraph.text == replace_str:
            table = document.add_table(rows=1, cols=cols_int)
            p = paragraph._p
            p.addprevious(table._tbl)
            hdr_cells = table.rows[0].cells
            for _title in xrange(cols_int):
                hdr_cells[_title].text =title_list[_title]

            for _index in xrange(len(info_list)):
                row_cells = table.add_row().cells
                for i in xrange(cols_int):
                    row_cells[i].text = info_list[_index][i]


    document.save(file_name_path)
if __name__ == '__main__':
    title = ['a','b','c','d']
    info = [["1","2","3","4"],["5","6","7","8"]]
    inert_table('demo11.docx',4,title,info,'{123}')
